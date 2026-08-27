from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("deliverables/QuickTalk_Internship_Final_Report.docx")
OUT.parent.mkdir(exist_ok=True)

NAVY = "0B2545"; BLUE = "2563EB"; TEAL = "0F766E"; PALE = "E8EEF5"
LIGHT = "F4F6F9"; MUTED = "5B6472"; WHITE = "FFFFFF"; RED = "9B1C1C"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
for name, size, color, before, after in [
    ("Title", 29, NAVY, 0, 8), ("Subtitle", 14, MUTED, 0, 12),
    ("Heading 1", 17, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, TEAL, 8, 4)]:
    s = styles[name]; s.font.name = "Calibri"; s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = name != "Subtitle"; s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=NAVY, size=9):
    cell.text = ""; p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(str(text)); r.bold = bold; r.font.name = "Calibri"; r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True, WHITE, 8.5); shade(t.rows[0].cells[i], NAVY)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, NAVY, 8.5)
            if ridx % 2: shade(cells[i], LIGHT)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths): row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4); p.add_run(text); return p

current_num_id = None
def _new_numbering_id():
    root = doc.part.numbering_part.element
    nums = root.findall(qn("w:num"))
    next_id = max([int(n.get(qn("w:numId"))) for n in nums] + [0]) + 1
    base = next((n for n in nums if n.get(qn("w:numId")) == "5"), nums[0])
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(next_id))
    abstract = OxmlElement("w:abstractNumId"); abstract.set(qn("w:val"), abstract_id); num.append(abstract)
    override = OxmlElement("w:lvlOverride"); override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride"); start.set(qn("w:val"), "1"); override.append(start); num.append(override)
    root.append(num); return next_id

def numbered(text, restart=False):
    global current_num_id
    if restart or current_num_id is None: current_num_id = _new_numbering_id()
    p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(4)
    numPr = p._p.get_or_add_pPr().get_or_add_numPr()
    numPr.get_or_add_ilvl().set(qn("w:val"), "0")
    numPr.get_or_add_numId().set(qn("w:val"), str(current_num_id))
    p.add_run(text); return p

def callout(label, text, color=BLUE):
    t = doc.add_table(rows=1, cols=1); t.autofit = False; t.columns[0].width = Inches(6.5)
    c = t.cell(0,0); shade(c, "EEF4FF")
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + "  "); r.bold = True; r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)

def page_break(): doc.add_page_break()

# Running furniture
h = sec.header.paragraphs[0]; h.text = "QUICKTALK  /  ORGANIZATIONAL MEMORY PLATFORM"; h.style = styles["Caption"]
h.runs[0].font.color.rgb = RGBColor.from_string(MUTED); h.runs[0].font.size = Pt(8)
f = sec.footer.paragraphs[0]; f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
f.add_run("Internship final report  |  ").font.size = Pt(8)
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); f._p.append(fld)

# Cover
doc.add_paragraph("INTERNSHIP FINAL REPORT", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("QuickTalk Organizational\nMemory & Agent Knowledge Platform", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Flask tool calling · Mem0 · Pinecone · Groq · Analytics", style="Subtitle"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("\n")
callout("Project result", "A working multi-tenant customer-support memory service that recalls prior conversations, creates concise agent handoffs, learns reusable human-agent knowledge safely, and exposes organization analytics.")
doc.add_paragraph("\n")
table(["Prepared for", "Prepared by", "Report date", "Repository status"], [["Internship supervisor / engineering lead", "Internship project owner", "27 August 2026", "Main branch · 48 tests passing"]], [1.7,1.7,1.3,1.8])
doc.add_paragraph("CONFIDENTIALITY NOTE", style="Heading 3")
doc.add_paragraph("API keys are intentionally excluded. Any key previously shared in chat must be rotated before production use.")

page_break()
doc.add_heading("1. Executive Summary", 1)
doc.add_paragraph("QuickTalk is a Flask-based support intelligence layer that gives a bot and human-agent team controlled access to organizational and customer conversation memory. The project connects live chat, cross-session recall, agent handoff, reusable human-agent knowledge, tone guidance, and analytics while keeping every lookup scoped to an organization and customer identity.")
table(["Capability", "Implemented outcome", "Primary technology"], [
    ["Conversational memory", "Cross-session semantic recall for a single customer", "Mem0 + Pinecone + Ollama embeddings"],
    ["Human handoff", "Three concise bullets and expandable session summaries", "Flask + precomputed profiles"],
    ["Knowledge fallback", "Bot KB first, approved agent knowledge second", "Versioned SQL state + Pinecone candidates"],
    ["Analytics", "Issues, sentiment, resolution and missing-topic trends", "Structured event store + deterministic aggregation"],
    ["Organization tone", "Style-only rewrite without changing facts", "Groq with safe fallback"],
], [1.35,3.25,1.9])
callout("Current status", "Locally demonstrable, committed to GitHub, real Mem0/Pinecone save-search verified, and 48 automated tests passing. It is a production-oriented prototype, not yet a production deployment.", TEAL)

doc.add_heading("2. Problem and Objectives", 1)
doc.add_paragraph("Customer-support conversations often restart from zero. Agents reread long histories, bots forget prior context, useful resolutions remain trapped in transcripts, and managers cannot see which topics the bot fails to answer. The project addresses these gaps through four objectives:")
for x in ["Remember useful customer context across sessions without mixing tenants or customers.", "Give a human agent a short, actionable handoff instead of a transcript wall.", "Turn only complete and reusable human-agent answers into governed organization knowledge.", "Convert conversation events into operational analytics and missing-knowledge priorities."]: bullet(x)

page_break()
doc.add_heading("3. System Architecture", 1)
doc.add_paragraph("The platform separates language generation, semantic retrieval, and authoritative business state. This is the central architectural decision.")
table(["Layer", "Responsibility", "Does not do"], [
    ["Browser UIs", "Chat, handoff, knowledge review, analytics and tool testing", "Store provider secrets"],
    ["Flask API", "Validate, authenticate, enforce scope and dispatch tools", "Act as a vector database"],
    ["Mem0", "Extract and retrieve conversational memories", "Own article lifecycle or dashboard totals"],
    ["Pinecone", "Store organization-namespaced vectors for semantic search", "Transactions, joins or authoritative version state"],
    ["Groq", "Summarization, curation and tone language tasks", "Persistent storage"],
    ["SQLite demo", "Messages, profiles, summaries, analytics, versions and audits", "Production-scale distributed operation"],
    ["Redis optional", "Shared low-latency profile cache in production", "System of record"],
], [1.2,3.1,2.2])
doc.add_heading("3.1 Request flow", 2)
for i, x in enumerate(["User action reaches a Flask page or JSON endpoint.", "Flask validates service key, required fields, role and organization scope.", "The tool registry chooses the correct memory, knowledge, handoff or analytics operation.", "Mem0/Pinecone or structured SQL retrieves evidence; Groq is called only where language generation is needed.", "The final response returns with grounding metadata and is stored as conversation history."]): numbered(x, restart=i == 0)
doc.add_heading("3.2 Tenant isolation", 2)
for x in ["Pinecone namespace is derived from organization scope.", "Customer Mem0 identity is an opaque hash of organization plus normalized mobile number.", "Relational queries always include organization scope.", "Flask rejects a mismatch between the claimed organization header and tool arguments.", "Production must derive organization and role from signed authentication claims."]: bullet(x)

page_break()
doc.add_heading("4. Tool Catalog and Activation", 1)
tools = [
 ("save_customer_memory","After every customer or assistant message","Writes conversation memory and analytics event","Mem0/Pinecone + SQL"),
 ("search_customer_memory","Explicit personal-history question","Cross-session semantic recall; general questions are skipped","Mem0/Pinecone + summary fallback"),
 ("get_customer_memory_context","New session initialization","Returns counts, current issue and recent summaries","Precomputed profile"),
 ("get_contextual_welcome","Chat opens / identity changes","Greets returning customer using latest useful issue","Profile + Groq tone"),
 ("get_handoff_context","Escalation to human","Returns three bullets, counts and expandable sessions","Profile + session summaries"),
 ("get_organization_tone_profile","Every customer-facing response / manual test","Returns style guidance only; no facts","Agent-message analysis"),
 ("get_knowledge_curation_policy","Agent chat closes / admin inspection","Returns tenant-specific save and reject rules","Versioned policy state"),
 ("search_approved_knowledge","Bot KB miss / manual test","Searches active current-version agent knowledge","Pinecone + SQL status gate"),
 ("resolve_support_answer","General support question","Bot KB → approved agent knowledge → apology","Knowledge orchestrator"),
 ("get_missing_knowledge_topics","Dashboard/report request","Aggregates unanswered topics without an LLM call","SQL aggregation"),
 ("import_agent_history_from_json","Explicit admin backfill","Starts controlled historical ingestion","Background importer"),
]
table(["Tool", "Activates when", "Outcome", "Core dependency"], tools, [1.55,1.55,2.15,1.25])

doc.add_heading("5. Detailed Tool Behavior", 1)
details = [
 ("save_customer_memory", "Inputs: organization, session, mobile, role, text, optional timestamp. The service normalizes identity, writes through Mem0/Pinecone, records a structured event, and refreshes the customer profile. Failure is surfaced without exposing provider keys."),
 ("search_customer_memory", "Runs only when the query explicitly refers to personal history. It searches across sessions, not only the current session. Results are combined with tenant/customer-scoped session summaries so incomplete vector backfills cannot break recall. The server decides answer eligibility."),
 ("get_customer_memory_context", "Loads a compact starting context for orchestration: memory count, previous-session count, current issue and recent summaries. It avoids sending the entire transcript to an LLM."),
 ("get_contextual_welcome", "Uses the latest meaningful issue to ask whether it was resolved. When no history exists, it returns a neutral greeting. Tone rewriting has a short timeout and cannot alter facts."),
 ("get_handoff_context", "Uses a precomputed profile for speed. The top of the inbox shows concise session summaries and counts; the full transcript loads only when the agent expands a session."),
 ("get_organization_tone_profile", "Learns reply-length, politeness and language-style signals. It explicitly marks facts_learned=false, keeping voice separate from knowledge."),
 ("get_knowledge_curation_policy", "Defines what can become reusable organizational knowledge. Controlled prices and regulations require scope, conditions and effective/version context."),
 ("search_approved_knowledge", "Pinecone supplies semantic candidates, but SQL verifies organization, active status and exact current article version. Disabled, deleted or superseded vectors are rejected."),
 ("resolve_support_answer", "Search order is bot knowledge first, approved agent knowledge second, safe apology third. A total miss creates a knowledge-gap event for analytics."),
 ("get_missing_knowledge_topics", "Groups questions that neither knowledge source could answer. Counts are deterministic database results; an LLM is not required for authoritative totals."),
 ("import_agent_history_from_json", "Administrative backfill for historical transcripts. It is not called during normal chat. Production should move this work to a queue with job IDs, idempotency and dead-letter handling."),
]
for name, text in details:
    doc.add_heading(name, 2); doc.add_paragraph(text)

page_break()
doc.add_heading("6. Main End-to-End Workflows", 1)
doc.add_heading("6.1 Live chat and RAG decision", 2)
for i, x in enumerate(["Customer asks a question.", "If it explicitly references personal history, search that customer’s Mem0/Pinecone memories across sessions.", "Otherwise search the organization’s bot knowledge base.", "If bot knowledge misses, search active approved human-agent knowledge.", "If both miss, apologize safely and record the missing topic.", "Apply organization tone to the grounded answer and store both messages."]): numbered(x, restart=i == 0)
callout("Important relevance rule", "A question such as “Is weekend installation available?” is a general policy question and must not be answered from unrelated personal memories. A question such as “What internet problem did I report previously?” is routed to personal memory.")
doc.add_heading("6.2 Human-agent handoff", 2)
for i, x in enumerate(["Agent clicks Escalate to human.", "Flask loads the precomputed customer profile and latest session summaries.", "The card displays concise history, session count and memory count.", "Older sessions and full transcripts remain expandable, reducing initial reading time."]): numbered(x, restart=i == 0)
doc.add_heading("6.3 Human-agent knowledge lifecycle", 2)
for i, x in enumerate(["Human-agent transcript is stored under the organization.", "When the chat closes, Groq proposes a standalone reusable question and answer.", "A deterministic validator rejects filler, referrals, incomplete answers and customer-specific data.", "Accepted content becomes an active versioned article and is indexed semantically.", "Editing creates a newer version; disabling or deleting prevents retrieval immediately through the SQL status gate."]): numbered(x, restart=i == 0)

doc.add_heading("7. Memory, Knowledge and Tone Boundaries", 1)
table(["Information type", "Correct home", "Example"], [
 ["Customer conversational memory", "Mem0/Pinecone", "Previously discussed doctor, issue or preference"],
 ["Reusable organization fact", "Versioned knowledge + Pinecone index", "Complete installation policy or current price"],
 ["Authoritative lifecycle state", "SQL", "Article active version, deletion and audit trail"],
 ["Analytics event", "SQL / warehouse", "Category, sentiment, resolution or gap"],
 ["Writing style", "Tone profile", "Concise, polite, Roman Urdu usage"],
], [1.55,2.05,2.9])

page_break()
doc.add_heading("8. Data Model and Indexing", 1)
table(["Table group", "Tables", "Purpose"], [
 ["Organizations and memory", "organizations, memory_events", "Tenant metadata and normalized conversation events"],
 ["Customer acceleration", "customer_profiles, session_summaries, durable_facts", "Fast handoff, welcome and recent-context access"],
 ["Gap analytics", "knowledge_gap_events", "Questions the bot could not answer"],
 ["Agent source records", "agent_sessions, agent_messages", "Original organization-owned human transcripts"],
 ["Governed knowledge", "knowledge_articles, knowledge_article_versions, bot_knowledge_articles", "Approved content, active versions and bot KB"],
 ["Governance", "knowledge_audit_log, knowledge_curation_policies", "Change history and tenant-specific rules"],
], [1.45,2.85,2.2])
doc.add_paragraph("Indexes cover organization plus time, session, mobile, category, status, gap state, durable-fact lookup and article/version status. These indexes support the dashboard and handoff without scanning full histories.")
doc.add_heading("8.1 Why Pinecone is not the only database", 2)
for x in ["Vector search is optimized for semantic similarity, not transactions and joins.", "Dashboards need exact grouped counts and time filters.", "Knowledge lifecycle needs authoritative active/deleted/version state.", "Auditability needs immutable source records and change logs.", "Production therefore uses Pinecone for retrieval and PostgreSQL/warehouse storage for authoritative business data."]: bullet(x)

doc.add_heading("9. Analytics", 1)
doc.add_paragraph("Every event is tenant- and time-scoped. The dashboard can show contact volume, categories, sentiment, resolution state, unresolved sessions and missing knowledge. Groq may help classify or summarize language; database aggregation remains authoritative.")
table(["Metric", "How it is produced", "Management use"], [
 ["Issue/category volume", "Grouped structured events", "Staffing and service priorities"],
 ["Resolution rate", "Session outcome/status", "Quality and escalation monitoring"],
 ["Sentiment", "Per-session/message classification", "Identify frustration and risk"],
 ["Missing topics", "Only total bot+agent-knowledge misses", "Plan new KB training"],
 ["Repeat contacts", "Customer/session counts", "Detect unresolved recurring problems"],
], [1.45,2.45,2.6])

page_break()
doc.add_heading("10. Security and Governance", 1)
for x in ["Provider and service keys remain server-side in an ignored .env file.", "Organization namespace and normalized customer identity are applied to every memory operation.", "Role checks protect human-agent and administrator actions.", "Knowledge is filtered through curation plus deterministic validation.", "Version/status verification prevents stale Pinecone vectors from being used.", "Keys exposed in messages must be rotated before production.", "Production needs signed auth claims, TLS, rate limiting, retention/deletion policy and secrets management."]: bullet(x)
doc.add_heading("11. Reliability and Performance", 1)
for x in ["Profiles and summaries are recomputed on new events rather than on every handoff.", "Local TTL cache accelerates the demo; Redis is the production shared-cache option.", "Mem0/Pinecone cold initialization has a bounded timeout; structured session summaries preserve continuity during provider delays.", "Tone has a short timeout and returns the original grounded answer if unavailable.", "SQL indexes support tenant/time and profile/session lookups.", "Production should add queue/outbox synchronization, retry limits, reconciliation and dead-letter handling."]: bullet(x)

doc.add_heading("12. Testing Evidence", 1)
table(["Evidence", "Result"], [
 ["Automated suite", "48 tests passing"],
 ["Real Pinecone index", "Ready, cosine metric, 768 dimensions"],
 ["Real Mem0 round trip", "Memory saved and retrieved from a different session"],
 ["Cross-session regression", "Current session no longer filters out earlier sessions"],
 ["Isolation", "Wrong organization/mobile cannot retrieve customer history"],
 ["Negative validation", "Bad keys, fields, roles, timestamps and scope mismatches covered"],
 ["Knowledge governance", "Rejection, version activation and disable/delete behavior covered"],
], [2.15,4.35])
callout("Verified example", "The test query “What was my previous verification code word?” returned the exact saved Mem0/Pinecone vector from an older session, alongside the structured-summary resilience source.", TEAL)

page_break()
doc.add_heading("13. Boss Demo Script", 1)
table(["Step", "Page", "Action", "What to explain"], [
 ["1", "/tools", "Show 11 callable schemas", "The service is tool-callable, not UI-only"],
 ["2", "/custom", "Ask a normal policy question", "Bot KB is searched before agent knowledge"],
 ["3", "/custom", "Ask: What internet problem did I report previously?", "Mem0/Pinecone cross-session recall"],
 ["4", "/custom", "Escalate to human", "Concise handoff with expandable sessions"],
 ["5", "/knowledge", "Open an agent transcript and article versions", "Strict reusable-knowledge governance"],
 ["6", "/dashboard", "Show categories, resolution, sentiment and missing topics", "Memory becomes operational insight"],
 ["7", "/api/health", "Show active backend", "mem0-groq-pinecone and cache status"],
], [0.45,0.8,2.4,2.85])
doc.add_heading("13.1 Suggested 60-second explanation", 2)
doc.add_paragraph("“I built a Flask tool-calling service that gives a support bot safe organizational memory. Mem0 orchestrates cross-session memory, Pinecone stores tenant-isolated vectors, Groq handles language tasks, and SQL keeps authoritative profiles, analytics and knowledge versions. The bot searches personal history only for personal-history questions; general questions use bot knowledge first and approved human-agent knowledge second. Human handoff shows concise session summaries, and unanswered topics become dashboard priorities. The implementation has 48 passing tests and a verified real Mem0/Pinecone round trip.”")

doc.add_heading("14. Production Roadmap", 1)
table(["Priority", "Change", "Reason"], [
 ["P0", "Rotate exposed keys and use a secrets manager", "Credential safety"],
 ["P0", "Signed authentication claims for tenant/user/role", "Prevent client-controlled scope"],
 ["P0", "PostgreSQL plus migrations and backups", "Concurrency, durability and scale"],
 ["P1", "Queue/outbox for Pinecone and imports", "Reliable synchronization and retries"],
 ["P1", "Redis shared cache", "Multi-worker performance"],
 ["P1", "Metrics, tracing, alerts and provider SLAs", "Operational visibility"],
 ["P2", "Retention, deletion and consent workflows", "Privacy and compliance"],
 ["P2", "Evaluation dataset and relevance monitoring", "Measure RAG quality over time"],
], [0.65,3.25,2.6])

page_break()
doc.add_heading("15. Limitations and Honest Status", 1)
for x in ["The current database is SQLite and is intended for demonstration, not multi-instance production.", "The Flask development server must be replaced by a production WSGI deployment.", "External providers can have cold starts, rate limits and outages; fallbacks reduce impact but do not remove dependency risk.", "Historical backfill quality depends on source transcript quality and strict curation.", "Analytics classification should be evaluated against labeled organization data before business decisions depend on it.", "Browser-side demo authentication must be replaced by real sessions and claims."]: bullet(x)
doc.add_heading("16. Internship Outcomes", 1)
for x in ["Designed a multi-tenant memory and retrieval architecture.", "Integrated Flask tool calling with Mem0, Pinecone, Groq and Ollama.", "Implemented relevance routing between personal memory, bot KB and agent knowledge.", "Built governed, versioned knowledge from human-agent transcripts.", "Created optimized handoff profiles and organization analytics.", "Diagnosed real integration failures, added resilience and regression tests.", "Documented production gaps rather than presenting a prototype as fully production-ready."]: bullet(x)

doc.add_heading("17. Glossary", 1)
table(["Term", "Plain-language meaning"], [
 ["RAG", "Retrieve relevant evidence before generating an answer"],
 ["Mem0", "Memory orchestration library for extracting and recalling conversational facts"],
 ["Embedding", "Numeric representation that allows search by meaning"],
 ["Pinecone namespace", "Organization-isolated vector partition"],
 ["Tenant", "One organization using the shared platform"],
 ["Handoff", "Context transferred from bot to human agent"],
 ["Knowledge gap", "A question neither approved knowledge source could answer"],
 ["Outbox", "Reliable event pattern for synchronizing database changes with external services"],
], [1.55,4.95])

doc.add_heading("Appendix A. Configuration Required", 1)
table(["Variable", "Purpose"], [
 ["SERVICE_API_KEY", "Protect tool and memory routes"], ["MEMORY_BACKEND", "Select mem0-groq"],
 ["PINECONE_API_KEY", "Authenticate vector storage"], ["MEM0_PINECONE_INDEX", "Select index"],
 ["GROQ_MEMORY_API_KEY", "Mem0 extraction"], ["GROQ_SUMMARY_API_KEY", "Session summaries and welcomes"],
 ["GROQ_KNOWLEDGE_API_KEY", "Knowledge curation and tone"], ["OLLAMA_BASE_URL", "Local embeddings"],
 ["MEM0_LOCAL_EMBEDDING_MODEL", "nomic-embed-text:latest"], ["ANALYTICS_DB_PATH", "Structured demo database"],
], [2.25,4.25])
doc.add_paragraph("No real credential value is included in this report.")

doc.add_heading("Appendix B. Main Demo URLs", 1)
for x in ["Customer chat and handoff: http://127.0.0.1:8765/custom", "Human-agent knowledge: http://127.0.0.1:8765/knowledge", "Organization analytics: http://127.0.0.1:8765/dashboard", "Tool catalog: http://127.0.0.1:8765/tools", "Health: http://127.0.0.1:8765/api/health"]: bullet(x)

doc.core_properties.title = "QuickTalk Organizational Memory Platform - Internship Final Report"
doc.core_properties.subject = "Architecture, tools, testing, demo and production roadmap"
doc.core_properties.author = "QuickTalk Internship Project"
doc.save(OUT)
print(OUT.resolve())
