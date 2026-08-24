from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).parent
SOURCE = ROOT / "IMPLEMENTATION_GUIDE.md"
OUTPUT = ROOT / "QuickTalk_Memory_Platform_Implementation_Guide.docx"
NAVY = RGBColor(24, 55, 88)
BLUE = RGBColor(44, 103, 160)
GRAY = RGBColor(90, 101, 114)


def font(run, size=10.5, bold=False, color=None, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    font(run, 8.5, color=GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`"):
            run = paragraph.add_run(part[1:-1]); font(run, 9.5, color=RGBColor(120, 45, 35), name="Consolas")
        elif part.startswith("**"):
            run = paragraph.add_run(part[2:-2]); font(run, bold=True)
        else:
            run = paragraph.add_run(part); font(run)


doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Inches(8.5), Inches(11)
section.top_margin = section.bottom_margin = Inches(0.78)
section.left_margin = section.right_margin = Inches(0.9)
section.header_distance, section.footer_distance = Inches(0.3), Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name, normal.font.size = "Aptos", Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.08
for name, size, color in (("Heading 1", 16, NAVY), ("Heading 2", 12.5, BLUE), ("Heading 3", 11, NAVY)):
    style = styles[name]
    style.font.name, style.font.size, style.font.bold, style.font.color.rgb = "Aptos Display", Pt(size), True, color
    style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(12), Pt(5)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
font(header.add_run("QUICKTALK  /  ENGINEERING IMPLEMENTATION GUIDE"), 8.5, bold=True, color=GRAY)
add_page_number(section.footer.paragraphs[0])

# Editorial-cover opening.
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(72); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("ENGINEERING REFERENCE"), 9, bold=True, color=BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
font(p.add_run("QuickTalk Memory and\nAgent Knowledge Platform"), 27, bold=True, color=NAVY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Flask • Mem0 • Pinecone • Groq • Analytics • Human Handoff"), 12, color=GRAY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(36)
font(p.add_run("Implementation architecture, tool APIs, examples, testing and production handoff"), 10.5, color=GRAY)
doc.add_page_break()

lines = SOURCE.read_text(encoding="utf-8").splitlines()[1:]
in_code = False
code_lines = []
for raw in lines:
    line = raw.rstrip()
    if line.startswith("```"):
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(5)
            run = p.add_run("\n".join(code_lines)); font(run, 8.2, color=RGBColor(45, 55, 65), name="Consolas")
            code_lines = []
        in_code = not in_code
        continue
    if in_code:
        code_lines.append(line); continue
    if not line:
        continue
    if line.startswith("## "):
        doc.add_heading(line[3:], level=1); continue
    if line.startswith("### "):
        doc.add_heading(line[4:], level=2); continue
    if re.match(r"^\d+\. ", line):
        p = doc.add_paragraph(style="List Number"); add_inline(p, re.sub(r"^\d+\. ", "", line)); continue
    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet"); add_inline(p, line[2:]); continue
    p = doc.add_paragraph(); add_inline(p, line)

doc.core_properties.title = "QuickTalk Memory and Agent Knowledge Platform"
doc.core_properties.subject = "Engineering implementation guide"
doc.core_properties.author = "QuickTalk Engineering"
doc.save(OUTPUT)
print(OUTPUT)
